# -*- encoding: utf-8 -*-

# @File        :   docx_parser.py
# @Time        :   2021/04/15 17:03:05
# @Author      :   
# @Email       :   
# @Description :   

import logging
import os
from typing import Dict, List
import zipfile
from lxml import etree

import docx
from docx2python.docx_context import get_context

from .base import BaseParser

logger = logging.getLogger("doc-extracter")


class Parser(BaseParser):

    supported_extension = '.docx'

    @classmethod
    def parse(cls, filename: str, method="xml") -> List[dict]:
        """ 解析 docx

        Args:
            filename (str): `.docx` 为扩展名的文件

        Raises:
            Exception: 不支持的文件类型/文件不存在

        Returns:
            List[dict]: 解析结果
        """

        if not os.path.exists(filename):
            raise Exception(f"Not Found: {filename}")

        if method == 'xml':
            return XMLParser.parse(filename)
        else:
            return CustomParser.parse(filename)


class XMLParser(object):

    w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    custom = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml'

    @classmethod
    def parse(cls, filename: str) -> List[dict]:
        f = zipfile.ZipFile(filename)
        context = get_context(f)
        # 读取主文档 xml
        document = context.get("officeDocument")
        # 读取文档关联 xml
        content_rels = (context.get("content_path2rels") or {}).get(document)

        # 读取主文档
        xml_content = f.read(document)
        tree = etree.fromstring(xml_content)

        section = []
        for index, paragraph in enumerate(cls._iterparagraphs(tree)):
            content = [text.strip() for text in cls._itertext(paragraph, f, content_rels) if text.strip()]
            if content:
                section.append({"page": str(index + 1), "context": " ".join(content)})
        return section

    @classmethod
    def _iterparagraphs(cls, my_etree):
        """Iterator to go through xml tree's text nodes"""
        for node in my_etree.iter(tag=etree.Element):
            if cls._check_element_is(node, "p"):
                yield node
            elif cls._check_element_is(node, "sdt"):
                if not cls._check_element_is(node.getparent(), "p"):
                    yield node

    @classmethod
    def _itertext(cls, my_etree, z: zipfile.ZipFile, content_rels: List[Dict]):
        """Iterator to go through xml tree's text nodes"""
        for node in my_etree.iter(tag=etree.Element):
            if cls._check_element_is(node, 't'):
                yield node.text
            elif cls._check_element_is(node, 'dataBinding'):
                yield cls._get_binding_data(node, z, content_rels)

    @classmethod
    def _iter_data_binding(cls, my_etree, z: zipfile.ZipFile, content_rels: List[Dict]):
        """Iterator to go through xml tree's dataBinding nodes"""
        for node in my_etree.iter(tag=etree.Element):
            if cls._check_element_is(node, 'dataBinding'):
                yield cls._get_binding_data(node, z, content_rels)

    @classmethod
    def _check_element_is(cls, element, type_char):
        return element.tag == '{%s}%s' % (cls.w, type_char)

    @classmethod
    def _get_binding_data(cls, element, z: zipfile.ZipFile, content_rels: List[Dict]):
        """Extract custom binding data"""
        if element.tag != '{%s}dataBinding' % cls.w:
            return ""
        
        xpath_key = '{%s}xpath' % cls.w
        xpath = element.attrib.get(xpath_key)

        mapping_key = '{%s}prefixMappings' % cls.w
        mapping_str = element.attrib.get(mapping_key)
        mappings = {}
        for item in mapping_str.split():
            tmp = item.split("=")
            if len(tmp) == 2:
                key, value  = tmp
                key = key.split(":")[-1]
                value = value.replace("'", "").replace('"', "")
                mappings[key] = value

        result = ""
        for rel in content_rels:
            if rel.get("Type") == cls.custom:
                try:
                    target = rel.get("Target")
                    target = "/".join([item for item in target.split("/") if item and item not in [".", ".."]])
                    node = etree.fromstring(z.read(target))
                    els = node.xpath(xpath, namespaces=mappings)
                    for el in els:
                        result += el.text
                    if els:
                        break
                except Exception as e:
                    logger.error(f"docx get data binding failed: {str(e)}")
                    continue
        return result


class CustomParser(object):

    @classmethod
    def parse(cls, filename: str) -> List[dict]:
        """ 解析 docx

        Args:
            filename (str): `.docx` 为扩展名的文件

        Raises:
            Exception: 不支持的文件类型/文件不存在

        Returns:
            List[dict]: 解析结果
        """

        if not os.path.exists(filename):
            raise Exception(f"Not Found: {filename}")

        section = []
        with open(filename, 'rb') as f:
            document = docx.Document(f)
            for i, paragraph in enumerate(document.paragraphs):
                if paragraph.text:
                    section.append({"page": str(i+1), "context": paragraph.text})

            # 处理表格
            table_data = []
            for table in document.tables:
                for row in table.rows:
                    table_data.extend([cell.text for cell in row.cells])
            if table_data:
                section.append({"page": "1", "context": "\n".join(table_data)})
        return section
